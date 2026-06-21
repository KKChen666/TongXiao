import os
import sys

_backend_dir = os.path.dirname(os.path.abspath(__file__))
if _backend_dir not in sys.path:
    sys.path.insert(0, _backend_dir)

from contextlib import asynccontextmanager
from fastapi import FastAPI, HTTPException, Depends, UploadFile, File, Form
from fastapi.middleware.cors import CORSMiddleware
from fastapi.staticfiles import StaticFiles
from pydantic import BaseModel
from database.db import (
    get_subjects, get_topics, get_cards,
    get_subject_progress, get_topic_progress, log_review, get_conn,
)
from database.migrate import run_migrations
from config import DB_TYPE
from auth import (
    hash_password, verify_password, create_token, get_current_user_id,
)
from ai.router import router as ai_router
from ai.admin import router as admin_router


@asynccontextmanager
async def lifespan(app: FastAPI):
    run_migrations()
    yield


app = FastAPI(title="考研背诵 API", lifespan=lifespan)
app.add_middleware(CORSMiddleware, allow_origins=["*"], allow_methods=["*"], allow_headers=["*"])
app.include_router(ai_router)
app.include_router(admin_router)


class RegisterBody(BaseModel):
    username: str
    password: str
    display_name: str = ""


class LoginBody(BaseModel):
    username: str
    password: str


class ReviewBody(BaseModel):
    result: int


class ImportBody(BaseModel):
    subject: str
    topic: str = ""
    cards: list


def _placeholder():
    return "%s" if DB_TYPE == "mysql" else "?"


# --- Auth Routes ---

@app.post("/api/register")
def api_register(body: RegisterBody):
    if len(body.username) < 2 or len(body.password) < 4:
        raise HTTPException(400, "用户名至少2位，密码至少4位")
    p = _placeholder()
    conn = get_conn()
    cur = conn.cursor()
    cur.execute(f"SELECT id FROM users WHERE username={p}", (body.username,))
    if cur.fetchone():
        raise HTTPException(400, "用户名已存在")
    display = body.display_name or body.username
    cur.execute(
        f"INSERT INTO users (username, password_hash, display_name) VALUES ({p},{p},{p})",
        (body.username, hash_password(body.password), display),
    )
    if DB_TYPE == "mysql":
        conn.commit()
    uid = cur.lastrowid
    token = create_token(uid, body.username)
    return {"token": token, "user": {"id": uid, "username": body.username, "display_name": display}}


@app.post("/api/login")
def api_login(body: LoginBody):
    p = _placeholder()
    cur = get_conn().cursor()
    cur.execute(f"SELECT id, username, password_hash, display_name FROM users WHERE username={p}", (body.username,))
    user = cur.fetchone()
    if not user or not verify_password(body.password, user["password_hash"]):
        raise HTTPException(401, "用户名或密码错误")
    token = create_token(user["id"], user["username"])
    return {"token": token, "user": {"id": user["id"], "username": user["username"], "display_name": user["display_name"]}}


@app.get("/api/me")
def api_me(user_id: int = Depends(get_current_user_id)):
    p = _placeholder()
    cur = get_conn().cursor()
    cur.execute(f"SELECT id, username, display_name, created_at FROM users WHERE id={p}", (user_id,))
    user = cur.fetchone()
    if not user:
        raise HTTPException(404, "用户不存在")
    return user


# --- Health Check (部署测试接口) ---

@app.get("/api/health")
def api_health():
    return {
        "status": "ok",
        "version": "v4-auto-deploy",
        "message": "TongXiao 后端自动部署验证成功！",
    }


# --- Public API Routes ---

@app.get("/api/subjects")
def api_subjects(user_id: int = Depends(get_current_user_id)):
    subs = get_subjects()
    result = []
    for s in subs:
        total, reviewed = get_subject_progress(s["id"], user_id)
        d = dict(s)
        d["total_cards"] = total
        d["reviewed_cards"] = reviewed
        d["progress_pct"] = round(reviewed / total * 100, 1) if total else 0
        result.append(d)
    return result


@app.get("/api/subjects/{subject_id}/topics")
def api_topics(subject_id: int, user_id: int = Depends(get_current_user_id)):
    tops = get_topics(subject_id)
    result = []
    for t in tops:
        total, reviewed = get_topic_progress(t["id"], user_id)
        d = dict(t)
        d["total_cards"] = total
        d["reviewed_cards"] = reviewed
        d["progress_pct"] = round(reviewed / total * 100, 1) if total else 0
        result.append(d)
    return result


@app.get("/api/topics/{topic_id}/cards")
def api_cards(topic_id: int):
    cards = get_cards(topic_id)
    if not cards:
        return []
    return [dict(c) for c in cards]


@app.post("/api/cards/{card_id}/review")
def api_review(card_id: int, body: ReviewBody, user_id: int = Depends(get_current_user_id)):
    log_review(card_id, body.result, user_id)
    return {"ok": True}


@app.get("/api/stats")
def api_stats(user_id: int = Depends(get_current_user_id)):
    subs = get_subjects()
    total_all = 0
    reviewed_all = 0
    for s in subs:
        total, reviewed = get_subject_progress(s["id"], user_id)
        total_all += total
        reviewed_all += reviewed
    subjects = []
    for s in subs:
        total, reviewed = get_subject_progress(s["id"], user_id)
        d = dict(s)
        d["total_cards"] = total
        d["reviewed_cards"] = reviewed
        d["progress_pct"] = round(reviewed / total * 100, 1) if total else 0
        subjects.append(d)
    return {
        "total_cards": total_all,
        "reviewed_cards": reviewed_all,
        "pending_cards": total_all - reviewed_all,
        "completion_rate": round(reviewed_all / total_all * 100, 1) if total_all else 0,
        "subjects": subjects,
    }


@app.post("/api/import")
def api_import(body: ImportBody, user_id: int = Depends(get_current_user_id)):
    if not body.cards:
        raise HTTPException(400, "No cards provided")
    conn = get_conn()
    cur = conn.cursor()
    p = _placeholder()
    cur.execute(f"SELECT id FROM subjects WHERE name={p}", (body.subject,))
    subj = cur.fetchone()
    if not subj:
        raise HTTPException(400, "Subject not found")
    sid = subj["id"]
    topic_name = body.topic or f"导入 - {body.subject}"
    cur.execute(f"SELECT id FROM topics WHERE subject_id={p} AND name={p}", (sid, topic_name))
    topic = cur.fetchone()
    if topic:
        tid = topic["id"]
        cur.execute(f"DELETE FROM cards WHERE topic_id={p}", (tid,))
    else:
        cur.execute(f"SELECT MAX(order_num) FROM topics WHERE subject_id={p}", (sid,))
        max_ord = cur.fetchone()["MAX(order_num)"] or 0
        cur.execute(f"INSERT INTO topics (subject_id, name, order_num) VALUES ({p},{p},{p})",
                    (sid, topic_name, max_ord + 1))
        tid = cur.lastrowid
    for idx, card in enumerate(body.cards):
        front = card.get("front", card.get("word", ""))
        back = card.get("back", card.get("definition", ""))
        cur.execute(f"INSERT INTO cards (topic_id, front, back, order_num) VALUES ({p},{p},{p},{p})",
                    (tid, front, back, idx + 1))
    if DB_TYPE == "mysql":
        conn.commit()
    return {"ok": True, "count": len(body.cards), "topic": topic_name}


@app.post("/api/import/file")
async def api_import_file(
    file: UploadFile = File(...),
    subject: str = Form("english"),
    topic: str = Form(""),
    user_id: int = Depends(get_current_user_id),
):
    filename = file.filename
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""

    content = await file.read()
    cards = []

    if ext == "csv":
        text = content.decode("utf-8", errors="replace")
        lines = text.split("\n")
        for line in lines[1:]:
            line = line.strip()
            if not line:
                continue
            parts = _parse_csv_line(line)
            if len(parts) >= 2 and parts[0] and parts[1]:
                cards.append({"front": parts[0], "back": parts[1]})

    elif ext == "json":
        import json
        text = content.decode("utf-8", errors="replace")
        parsed = json.loads(text)
        if isinstance(parsed, list):
            for r in parsed:
                f = r.get("front") or r.get("word") or ""
                b = r.get("back") or r.get("definition") or r.get("meaning") or ""
                if f and b:
                    cards.append({"front": f, "back": b})

    elif ext == "txt":
        text = content.decode("utf-8", errors="replace")
        for line in text.split("\n"):
            line = line.strip()
            if not line:
                continue
            parts = line.split("|") if "|" in line else line.split("\t")
            if len(parts) >= 2 and parts[0].strip() and parts[1].strip():
                cards.append({"front": parts[0].strip(), "back": parts[1].strip()})

    elif ext == "docx":
        try:
            import docx
            import io
            doc = docx.Document(io.BytesIO(content))
            paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
            cards = _extract_cards_from_lines(paragraphs)
        except ImportError:
            raise HTTPException(400, "服务器未安装 python-docx，无法解析 Word 文件")
        except Exception as e:
            raise HTTPException(400, f"Word 文件解析失败: {e}")

    elif ext == "pdf":
        try:
            import pdfplumber
            import io
            text_parts = []
            with pdfplumber.open(io.BytesIO(content)) as pdf:
                for page in pdf.pages:
                    t = page.extract_text()
                    if t:
                        text_parts.append(t)
            full_text = "\n".join(text_parts)
            lines = [l.strip() for l in full_text.split("\n") if l.strip()]
            cards = _extract_cards_from_lines(lines)
        except ImportError:
            raise HTTPException(400, "服务器未安装 pdfplumber，无法解析 PDF 文件")
        except Exception as e:
            raise HTTPException(400, f"PDF 文件解析失败: {e}")

    else:
        raise HTTPException(400, f"不支持的文件格式: .{ext}")

    if not cards:
        raise HTTPException(400, "未能从文件中提取到有效的卡片数据")

    conn = get_conn()
    cur = conn.cursor()
    p = _placeholder()
    cur.execute(f"SELECT id FROM subjects WHERE name={p}", (subject,))
    subj = cur.fetchone()
    if not subj:
        raise HTTPException(400, "Subject not found")
    sid = subj["id"]
    topic_name = topic or f"导入 - {filename}"
    cur.execute(f"SELECT id FROM topics WHERE subject_id={p} AND name={p}", (sid, topic_name))
    existing = cur.fetchone()
    if existing:
        tid = existing["id"]
        cur.execute(f"DELETE FROM cards WHERE topic_id={p}", (tid,))
    else:
        cur.execute(f"SELECT MAX(order_num) FROM topics WHERE subject_id={p}", (sid,))
        max_ord = cur.fetchone()["MAX(order_num)"] or 0
        cur.execute(f"INSERT INTO topics (subject_id, name, order_num) VALUES ({p},{p},{p})",
                    (sid, topic_name, max_ord + 1))
        tid = cur.lastrowid
    for idx, card in enumerate(cards):
        cur.execute(f"INSERT INTO cards (topic_id, front, back, order_num) VALUES ({p},{p},{p},{p})",
                    (tid, card["front"], card["back"], idx + 1))
    if DB_TYPE == "mysql":
        conn.commit()
    return {"ok": True, "count": len(cards), "topic": topic_name}


def _parse_csv_line(line):
    result = []
    current = ""
    in_quotes = False
    for ch in line:
        if in_quotes:
            if ch == '"':
                in_quotes = False
            else:
                current += ch
        else:
            if ch == '"':
                in_quotes = True
            elif ch == ',':
                result.append(current.strip())
                current = ""
            else:
                current += ch
    result.append(current.strip())
    return result


def _extract_cards_from_lines(lines):
    """Try to extract front/back card pairs from lines of text.
    Supports patterns:
      - word\n  definition  (two consecutive lines)
      - word: definition
      - word - definition
      - word\tdefinition
    """
    cards = []

    # Pattern 1: lines with separator (: or - or tab) on same line
    for line in lines:
        for sep in ["\t", "：", ": ", " - ", " — "]:
            if sep in line:
                parts = line.split(sep, 1)
                if len(parts) == 2 and parts[0].strip() and parts[1].strip():
                    cards.append({"front": parts[0].strip(), "back": parts[1].strip()})
                    break
    if cards:
        return cards

    # Pattern 2: consecutive line pairs (odd = front, even = back)
    if len(lines) >= 2:
        i = 0
        while i < len(lines) - 1:
            front = lines[i].strip()
            back = lines[i + 1].strip()
            # Heuristic: if front line is short (<80 chars) and back is longer or different
            if front and back and len(front) < 120:
                cards.append({"front": front, "back": back})
                i += 2
            else:
                i += 1
    return cards


def _compute_hash(front: str, subject_id: int) -> str:
    import hashlib
    raw = f"{front.strip().lower()}|{subject_id}"
    return hashlib.sha256(raw.encode("utf-8")).hexdigest()[:64]


# --- Knowledge Base API Routes ---

class KnowledgeImportBody(BaseModel):
    subject: str
    tags: list = []
    items: list


@app.get("/api/knowledge")
def api_knowledge(
    subject: str = "english",
    search: str = "",
    tag: str = "",
    page: int = 1,
    page_size: int = 50,
    user_id: int = Depends(get_current_user_id),
):
    p = _placeholder()
    conn = get_conn()
    cur = conn.cursor()

    cur.execute(f"SELECT id FROM subjects WHERE name={p}", (subject,))
    subj = cur.fetchone()
    if not subj:
        raise HTTPException(400, "Subject not found")
    sid = subj["id"]

    conditions = [f"subject_id={p}"]
    params = [sid]

    if search:
        conditions.append(f"(front LIKE {p} OR back LIKE {p})")
        params.extend([f"%{search}%", f"%{search}%"])

    if tag:
        conditions.append(f"tags LIKE {p}")
        params.append(f'%"{tag}"%')

    where = " AND ".join(conditions)

    cur.execute(f"SELECT COUNT(*) as cnt FROM knowledge_items WHERE {where}", tuple(params))
    total = cur.fetchone()["cnt"]

    offset = (page - 1) * page_size
    cur.execute(
        f"SELECT * FROM knowledge_items WHERE {where} ORDER BY id LIMIT {p} OFFSET {p}",
        tuple(params) + (page_size, offset),
    )
    items = [dict(r) for r in cur.fetchall()]

    return {
        "items": items,
        "total": total,
        "page": page,
        "page_size": page_size,
        "total_pages": (total + page_size - 1) // page_size,
    }


@app.get("/api/knowledge/tags")
def api_knowledge_tags(subject: str = "english", user_id: int = Depends(get_current_user_id)):
    p = _placeholder()
    conn = get_conn()
    cur = conn.cursor()

    cur.execute(f"SELECT id FROM subjects WHERE name={p}", (subject,))
    subj = cur.fetchone()
    if not subj:
        raise HTTPException(400, "Subject not found")
    sid = subj["id"]

    cur.execute(f"SELECT tags FROM knowledge_items WHERE subject_id={p}", (sid,))
    rows = cur.fetchall()

    tag_counts = {}
    for row in rows:
        import json
        try:
            tags = json.loads(row["tags"]) if row["tags"] else []
            for t in tags:
                tag_counts[t] = tag_counts.get(t, 0) + 1
        except:
            pass

    return [{"name": k, "count": v} for k, v in sorted(tag_counts.items())]


@app.post("/api/knowledge/import")
def api_knowledge_import(body: KnowledgeImportBody, user_id: int = Depends(get_current_user_id)):
    if not body.items:
        raise HTTPException(400, "No items provided")

    p = _placeholder()
    conn = get_conn()
    cur = conn.cursor()

    cur.execute(f"SELECT id FROM subjects WHERE name={p}", (body.subject,))
    subj = cur.fetchone()
    if not subj:
        raise HTTPException(400, "Subject not found")
    sid = subj["id"]

    import json
    tags_json = json.dumps(body.tags, ensure_ascii=False) if body.tags else "[]"

    new_count = 0
    skip_count = 0

    for item in body.items:
        front = item.get("front", item.get("word", ""))
        back = item.get("back", item.get("definition", ""))
        if not front or not back:
            skip_count += 1
            continue

        h = _compute_hash(front, sid)

        if DB_TYPE == "mysql":
            sql = (
                f"INSERT IGNORE INTO knowledge_items "
                f"(subject_id, front, back, tags, hash, source, created_by) "
                f"VALUES ({p}, {p}, {p}, {p}, {p}, 'import', {p})"
            )
        else:
            sql = (
                f"INSERT OR IGNORE INTO knowledge_items "
                f"(subject_id, front, back, tags, hash, source, created_by) "
                f"VALUES ({p}, {p}, {p}, {p}, {p}, 'import', {p})"
            )

        cur.execute(sql, (sid, front, back, tags_json, h, user_id))

        if cur.rowcount > 0:
            new_count += 1
        else:
            skip_count += 1

    if DB_TYPE == "mysql":
        conn.commit()

    return {"ok": True, "new": new_count, "duplicate": skip_count}


@app.post("/api/knowledge/import/file")
async def api_knowledge_import_file(
    file: UploadFile = File(...),
    subject: str = Form("english"),
    tags: str = Form(""),
    user_id: int = Depends(get_current_user_id),
):
    filename = file.filename
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""

    content = await file.read()
    items = []

    if ext == "csv":
        text = content.decode("utf-8", errors="replace")
        lines = text.split("\n")
        for line in lines[1:]:
            line = line.strip()
            if not line:
                continue
            parts = _parse_csv_line(line)
            if len(parts) >= 2 and parts[0] and parts[1]:
                items.append({"front": parts[0], "back": parts[1]})

    elif ext == "json":
        import json as json_mod
        text = content.decode("utf-8", errors="replace")
        parsed = json_mod.loads(text)
        if isinstance(parsed, list):
            for r in parsed:
                f = r.get("front") or r.get("word") or ""
                b = r.get("back") or r.get("definition") or r.get("meaning") or ""
                if f and b:
                    items.append({"front": f, "back": b})

    elif ext == "txt":
        text = content.decode("utf-8", errors="replace")
        for line in text.split("\n"):
            line = line.strip()
            if not line:
                continue
            parts = line.split("|") if "|" in line else line.split("\t")
            if len(parts) >= 2 and parts[0].strip() and parts[1].strip():
                items.append({"front": parts[0].strip(), "back": parts[1].strip()})

    elif ext == "docx":
        try:
            import docx
            import io
            doc = docx.Document(io.BytesIO(content))
            paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
            items = _extract_cards_from_lines(paragraphs)
        except ImportError:
            raise HTTPException(400, "服务器未安装 python-docx，无法解析 Word 文件")
        except Exception as e:
            raise HTTPException(400, f"Word 文件解析失败: {e}")

    elif ext == "pdf":
        try:
            import pdfplumber
            import io
            text_parts = []
            with pdfplumber.open(io.BytesIO(content)) as pdf:
                for page in pdf.pages:
                    t = page.extract_text()
                    if t:
                        text_parts.append(t)
            full_text = "\n".join(text_parts)
            lines = [l.strip() for l in full_text.split("\n") if l.strip()]
            items = _extract_cards_from_lines(lines)
        except ImportError:
            raise HTTPException(400, "服务器未安装 pdfplumber，无法解析 PDF 文件")
        except Exception as e:
            raise HTTPException(400, f"PDF 文件解析失败: {e}")

    else:
        raise HTTPException(400, f"不支持的文件格式: .{ext}")

    if not items:
        raise HTTPException(400, "未能从文件中提取到有效的词条数据")

    import json as json_mod
    tag_list = [t.strip() for t in tags.split(",") if t.strip()] if tags else []

    p = _placeholder()
    conn = get_conn()
    cur = conn.cursor()

    cur.execute(f"SELECT id FROM subjects WHERE name={p}", (subject,))
    subj = cur.fetchone()
    if not subj:
        raise HTTPException(400, "Subject not found")
    sid = subj["id"]

    tags_json = json_mod.dumps(tag_list, ensure_ascii=False) if tag_list else "[]"

    new_count = 0
    skip_count = 0

    for item in items:
        h = _compute_hash(item["front"], sid)

        if DB_TYPE == "mysql":
            sql = (
                f"INSERT IGNORE INTO knowledge_items "
                f"(subject_id, front, back, tags, hash, source, created_by) "
                f"VALUES ({p}, {p}, {p}, {p}, {p}, 'import', {p})"
            )
        else:
            sql = (
                f"INSERT OR IGNORE INTO knowledge_items "
                f"(subject_id, front, back, tags, hash, source, created_by) "
                f"VALUES ({p}, {p}, {p}, {p}, {p}, 'import', {p})"
            )

        cur.execute(sql, (sid, item["front"], item["back"], tags_json, h, user_id))

        if cur.rowcount > 0:
            new_count += 1
        else:
            skip_count += 1

    if DB_TYPE == "mysql":
        conn.commit()

    return {"ok": True, "new": new_count, "duplicate": skip_count, "filename": filename}


# --- Wordbook API Routes ---

class WordbookCreateBody(BaseModel):
    name: str
    description: str = ""
    subject: str
    item_ids: list = []


@app.get("/api/wordbooks")
def api_wordbooks(user_id: int = Depends(get_current_user_id)):
    p = _placeholder()
    conn = get_conn()
    cur = conn.cursor()

    cur.execute(
        f"SELECT w.*, s.display_name as subject_name FROM wordbooks w "
        f"JOIN subjects s ON w.subject_id = s.id "
        f"WHERE w.user_id={p} ORDER BY w.updated_at DESC",
        (user_id,),
    )
    return [dict(r) for r in cur.fetchall()]


@app.post("/api/wordbooks")
def api_wordbook_create(body: WordbookCreateBody, user_id: int = Depends(get_current_user_id)):
    p = _placeholder()
    conn = get_conn()
    cur = conn.cursor()

    cur.execute(f"SELECT id FROM subjects WHERE name={p}", (body.subject,))
    subj = cur.fetchone()
    if not subj:
        raise HTTPException(400, "Subject not found")
    sid = subj["id"]

    cur.execute(
        f"INSERT INTO wordbooks (user_id, name, description, subject_id) VALUES ({p},{p},{p},{p})",
        (user_id, body.name, body.description, sid),
    )
    if DB_TYPE == "mysql":
        conn.commit()
    wb_id = cur.lastrowid

    added = 0
    if body.item_ids:
        for idx, kid in enumerate(body.item_ids):
            try:
                cur.execute(
                    f"INSERT INTO wordbook_items (wordbook_id, knowledge_id, order_num) VALUES ({p},{p},{p})",
                    (wb_id, kid, idx + 1),
                )
                added += 1
            except:
                pass
        cur.execute(f"UPDATE wordbooks SET item_count={p} WHERE id={p}", (added, wb_id))
        if DB_TYPE == "mysql":
            conn.commit()

    return {"ok": True, "id": wb_id, "item_count": added}


@app.get("/api/wordbooks/{wordbook_id}")
def api_wordbook_detail(wordbook_id: int, user_id: int = Depends(get_current_user_id)):
    p = _placeholder()
    conn = get_conn()
    cur = conn.cursor()

    cur.execute(
        f"SELECT w.*, s.display_name as subject_name, s.name as subject_key FROM wordbooks w "
        f"JOIN subjects s ON w.subject_id = s.id "
        f"WHERE w.id={p} AND w.user_id={p}",
        (wordbook_id, user_id),
    )
    wb = cur.fetchone()
    if not wb:
        raise HTTPException(404, "词书不存在")

    cur.execute(
        f"SELECT wi.id as item_id, wi.order_num, k.* FROM wordbook_items wi "
        f"JOIN knowledge_items k ON wi.knowledge_id = k.id "
        f"WHERE wi.wordbook_id={p} ORDER BY wi.order_num",
        (wordbook_id,),
    )
    items = [dict(r) for r in cur.fetchall()]

    result = dict(wb)
    result["items"] = items
    return result


@app.post("/api/wordbooks/{wordbook_id}/items")
def api_wordbook_add_items(wordbook_id: int, body: dict, user_id: int = Depends(get_current_user_id)):
    p = _placeholder()
    conn = get_conn()
    cur = conn.cursor()

    cur.execute(f"SELECT id FROM wordbooks WHERE id={p} AND user_id={p}", (wordbook_id, user_id))
    if not cur.fetchone():
        raise HTTPException(404, "词书不存在")

    item_ids = body.get("item_ids", [])
    if not item_ids:
        raise HTTPException(400, "No item_ids provided")

    cur.execute(f"SELECT MAX(order_num) as max_ord FROM wordbook_items WHERE wordbook_id={p}", (wordbook_id,))
    max_ord = cur.fetchone()["max_ord"] or 0

    added = 0
    for idx, kid in enumerate(item_ids):
        try:
            cur.execute(
                f"INSERT INTO wordbook_items (wordbook_id, knowledge_id, order_num) VALUES ({p},{p},{p})",
                (wordbook_id, kid, max_ord + idx + 1),
            )
            added += 1
        except:
            pass

    cur.execute(
        f"UPDATE wordbooks SET item_count=(SELECT COUNT(*) FROM wordbook_items WHERE wordbook_id={p}) WHERE id={p}",
        (wordbook_id, wordbook_id),
    )
    if DB_TYPE == "mysql":
        conn.commit()

    return {"ok": True, "added": added}


@app.delete("/api/wordbooks/{wordbook_id}/items/{knowledge_id}")
def api_wordbook_remove_item(wordbook_id: int, knowledge_id: int, user_id: int = Depends(get_current_user_id)):
    p = _placeholder()
    conn = get_conn()
    cur = conn.cursor()

    cur.execute(f"SELECT id FROM wordbooks WHERE id={p} AND user_id={p}", (wordbook_id, user_id))
    if not cur.fetchone():
        raise HTTPException(404, "词书不存在")

    cur.execute(
        f"DELETE FROM wordbook_items WHERE wordbook_id={p} AND knowledge_id={p}",
        (wordbook_id, knowledge_id),
    )

    cur.execute(
        f"UPDATE wordbooks SET item_count=(SELECT COUNT(*) FROM wordbook_items WHERE wordbook_id={p}) WHERE id={p}",
        (wordbook_id, wordbook_id),
    )
    if DB_TYPE == "mysql":
        conn.commit()

    return {"ok": True}


@app.delete("/api/wordbooks/{wordbook_id}")
def api_wordbook_delete(wordbook_id: int, user_id: int = Depends(get_current_user_id)):
    p = _placeholder()
    conn = get_conn()
    cur = conn.cursor()

    cur.execute(f"SELECT id FROM wordbooks WHERE id={p} AND user_id={p}", (wordbook_id, user_id))
    if not cur.fetchone():
        raise HTTPException(404, "词书不存在")

    cur.execute(f"DELETE FROM wordbooks WHERE id={p}", (wordbook_id,))
    if DB_TYPE == "mysql":
        conn.commit()

    return {"ok": True}


@app.post("/api/wordbooks/{wordbook_id}/to-topic")
def api_wordbook_to_topic(wordbook_id: int, user_id: int = Depends(get_current_user_id)):
    p = _placeholder()
    conn = get_conn()
    cur = conn.cursor()

    cur.execute(
        f"SELECT w.*, s.name as subject_key FROM wordbooks w "
        f"JOIN subjects s ON w.subject_id = s.id "
        f"WHERE w.id={p} AND w.user_id={p}",
        (wordbook_id, user_id),
    )
    wb = cur.fetchone()
    if not wb:
        raise HTTPException(404, "词书不存在")

    cur.execute(
        f"SELECT k.front, k.back FROM wordbook_items wi "
        f"JOIN knowledge_items k ON wi.knowledge_id = k.id "
        f"WHERE wi.wordbook_id={p} ORDER BY wi.order_num",
        (wordbook_id,),
    )
    items = cur.fetchall()

    if not items:
        raise HTTPException(400, "词书为空，无法生成学习卡片")

    sid = wb["subject_id"]
    topic_name = wb["name"]

    cur.execute(f"SELECT id FROM topics WHERE subject_id={p} AND name={p}", (sid, topic_name))
    existing = cur.fetchone()
    if existing:
        tid = existing["id"]
        cur.execute(f"DELETE FROM cards WHERE topic_id={p}", (tid,))
    else:
        cur.execute(f"SELECT MAX(order_num) FROM topics WHERE subject_id={p}", (sid,))
        max_ord = cur.fetchone()["MAX(order_num)"] or 0
        cur.execute(
            f"INSERT INTO topics (subject_id, name, order_num, book_id) VALUES ({p},{p},{p}, NULL)",
            (sid, topic_name, max_ord + 1),
        )
        tid = cur.lastrowid

    for idx, item in enumerate(items):
        cur.execute(
            f"INSERT INTO cards (topic_id, front, back, order_num) VALUES ({p},{p},{p},{p})",
            (tid, item["front"], item["back"], idx + 1),
        )

    if DB_TYPE == "mysql":
        conn.commit()

    return {"ok": True, "topic_id": tid, "topic_name": topic_name, "card_count": len(items)}


class UnitTopicBody(BaseModel):
    unit_name: str
    offset: int = 0
    limit: int = 20


@app.post("/api/wordbooks/{wordbook_id}/unit-to-topic")
def api_wordbook_unit_to_topic(wordbook_id: int, body: UnitTopicBody, user_id: int = Depends(get_current_user_id)):
    p = _placeholder()
    conn = get_conn()
    cur = conn.cursor()

    cur.execute(
        f"SELECT w.*, s.name as subject_key FROM wordbooks w "
        f"JOIN subjects s ON w.subject_id = s.id "
        f"WHERE w.id={p} AND w.user_id={p}",
        (wordbook_id, user_id),
    )
    wb = cur.fetchone()
    if not wb:
        raise HTTPException(404, "词书不存在")

    cur.execute(
        f"SELECT k.front, k.back FROM wordbook_items wi "
        f"JOIN knowledge_items k ON wi.knowledge_id = k.id "
        f"WHERE wi.wordbook_id={p} ORDER BY wi.order_num LIMIT {p} OFFSET {p}",
        (wordbook_id, body.limit, body.offset),
    )
    items = cur.fetchall()

    if not items:
        raise HTTPException(400, "该单元没有词条")

    sid = wb["subject_id"]
    topic_name = f"{wb['name']} - {body.unit_name}"

    cur.execute(f"SELECT id FROM topics WHERE subject_id={p} AND name={p}", (sid, topic_name))
    existing = cur.fetchone()
    if existing:
        tid = existing["id"]
        cur.execute(f"DELETE FROM cards WHERE topic_id={p}", (tid,))
    else:
        cur.execute(f"SELECT MAX(order_num) FROM topics WHERE subject_id={p}", (sid,))
        max_ord = cur.fetchone()["MAX(order_num)"] or 0
        cur.execute(
            f"INSERT INTO topics (subject_id, name, order_num, book_id) VALUES ({p},{p},{p}, NULL)",
            (sid, topic_name, max_ord + 1),
        )
        tid = cur.lastrowid

    for idx, item in enumerate(items):
        cur.execute(
            f"INSERT INTO cards (topic_id, front, back, order_num) VALUES ({p},{p},{p},{p})",
            (tid, item["front"], item["back"], idx + 1),
        )

    if DB_TYPE == "mysql":
        conn.commit()

    return {"ok": True, "topic_id": tid, "topic_name": topic_name, "card_count": len(items)}


# --- Serve Static Frontend (production build) ---

STATIC_DIR = os.path.join(os.path.dirname(__file__), "frontend", "dist")
if not os.path.isdir(STATIC_DIR):
    STATIC_DIR = os.path.join(os.path.dirname(__file__), "..", "frontend", "dist")
if os.path.isdir(STATIC_DIR):
    app.mount("/", StaticFiles(directory=STATIC_DIR, html=True), name="static")
